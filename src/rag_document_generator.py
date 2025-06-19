"""
Document generator for RAG tutorial demonstrations.
Creates sample documents about Strands framework for educational purposes.
"""

import os
from pathlib import Path
from datetime import datetime
from typing import List, Dict

def ensure_rag_docs_dir():
    """Ensure the rag_docs directory exists."""
    rag_docs_path = Path("rag_docs")
    rag_docs_path.mkdir(exist_ok=True)
    return rag_docs_path

def create_basic_documents():
    """Create simple text documents for the basic RAG tutorial."""
    rag_docs_path = ensure_rag_docs_dir()
    
    documents = {
        "strands_framework_guide.md": """# Strands Framework Guide

## What is Strands?

Strands is a powerful Python framework for building AI agents. It provides a simple yet flexible way to create intelligent assistants that can understand natural language, use tools, and maintain conversation context.

## Key Features

- **Easy Agent Creation**: Create agents with just a few lines of code
- **Multiple Model Support**: Use AWS Bedrock, OpenAI, Anthropic, or local models
- **Tool Integration**: Give agents abilities through custom tools
- **Conversation Memory**: Agents remember context across interactions
- **Streaming Support**: Real-time responses for better user experience

## Basic Usage

```python
from strands import Agent
agent = Agent(system_prompt="You are a helpful assistant")
response = agent("Hello!")
```

## Agent Configuration

Agents can be customized with:
- System prompts to define behavior
- Tools to extend capabilities
- Different language models
- Memory systems for context retention

## Best Practices

1. Write clear system prompts
2. Use type hints in tools
3. Handle errors gracefully
4. Test with different inputs
5. Monitor performance
""",

        "python_best_practices.txt": """Python Best Practices for AI Development

1. Code Organization
   - Use clear, descriptive variable names
   - Follow PEP 8 style guidelines
   - Organize code into logical modules
   - Keep functions small and focused

2. Type Hints
   - Always use type hints for function parameters
   - Specify return types
   - Use Optional for nullable values
   - Import from typing module

3. Error Handling
   - Use try-except blocks appropriately
   - Log errors for debugging
   - Provide meaningful error messages
   - Don't silence exceptions

4. Documentation
   - Write clear docstrings
   - Include examples in documentation
   - Document edge cases
   - Keep README files updated

5. Testing
   - Write unit tests for tools
   - Test edge cases
   - Use pytest for testing
   - Aim for high test coverage

6. Performance
   - Profile code for bottlenecks
   - Use generators for large datasets
   - Cache expensive computations
   - Optimize database queries

7. Security
   - Never hardcode credentials
   - Use environment variables
   - Validate user inputs
   - Follow OWASP guidelines
""",

        "ai_agents_faq.md": """# AI Agents Frequently Asked Questions

## What is an AI agent?

An AI agent is a software program that can perceive its environment, make decisions, and take actions to achieve specific goals. In the context of Strands, agents are powered by large language models and can interact with users through natural language.

## How do agents differ from chatbots?

While chatbots typically follow scripted responses, AI agents can:
- Reason about complex problems
- Use tools to perform actions
- Maintain context over long conversations
- Adapt their behavior based on instructions

## What are tools in the context of agents?

Tools are functions that agents can call to perform specific tasks. Examples include:
- Searching databases
- Performing calculations
- Accessing APIs
- Reading files
- Sending emails

## How do I choose the right model?

Consider these factors:
- **Task complexity**: Complex reasoning needs powerful models
- **Speed requirements**: Smaller models respond faster
- **Cost**: Cloud models charge per token
- **Privacy**: Local models keep data on-device

## Can agents learn from conversations?

Agents maintain conversation history within a session but don't permanently learn. For persistent learning, you need to implement custom memory systems or fine-tune models.

## What are system prompts?

System prompts define an agent's personality, behavior, and capabilities. They're instructions that guide how the agent responds to user inputs.

## How do I handle errors in agents?

- Implement proper error handling in tools
- Provide fallback responses
- Log errors for debugging
- Test edge cases thoroughly
""",

        "getting_started_bedrock.txt": """Getting Started with AWS Bedrock

Prerequisites:
- AWS Account with Bedrock access
- AWS CLI installed and configured
- Python 3.8 or higher
- Strands SDK installed

Step 1: Enable Model Access
1. Log into AWS Console
2. Navigate to Amazon Bedrock
3. Go to "Model access"
4. Request access to Claude 3.7 Sonnet
5. Wait for approval (usually instant)

Step 2: Configure AWS Credentials
Option A - AWS CLI:
$ aws configure
Enter your access key, secret key, and region

Option B - Environment Variables:
export AWS_ACCESS_KEY_ID="your-key"
export AWS_SECRET_ACCESS_KEY="your-secret"
export AWS_DEFAULT_REGION="us-west-2"

Step 3: Test Bedrock Connection
from strands.models import BedrockModel
model = BedrockModel(model_id="claude-3.7-sonnet")

Step 4: Create Your First Agent
from strands import Agent
agent = Agent(model=model)
response = agent("Hello!")

Common Issues:
- "Access Denied": Check model access in Bedrock console
- "Invalid credentials": Verify AWS configuration
- "Region error": Ensure Bedrock is available in your region

Tips:
- Start with claude-3.7-sonnet for best results
- Use IAM roles in production
- Monitor costs in AWS Cost Explorer
- Set up CloudWatch for logging
"""
    }
    
    # Write all documents
    for filename, content in documents.items():
        filepath = rag_docs_path / filename
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
    
    return list(documents.keys())

def create_advanced_documents():
    """Create more complex documents for the advanced RAG tutorial."""
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.units import inch
    
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        DOCX_AVAILABLE = True
    except ImportError:
        DOCX_AVAILABLE = False
        print("Warning: python-docx not installed. Skipping .docx file generation.")
    
    rag_docs_path = ensure_rag_docs_dir()
    
    # Create PDF - API Reference
    pdf_path = rag_docs_path / "strands_api_reference.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    width, height = letter
    
    # Page 1 - Title and Overview
    c.setFont("Helvetica-Bold", 24)
    c.drawString(100, height - 100, "Strands API Reference")
    c.setFont("Helvetica", 12)
    c.drawString(100, height - 140, "Version 0.1.x")
    c.drawString(100, height - 160, f"Generated: {datetime.now().strftime('%Y-%m-%d')}")
    
    c.setFont("Helvetica-Bold", 16)
    c.drawString(100, height - 220, "Agent Class")
    c.setFont("Helvetica", 11)
    y_pos = height - 250
    
    agent_docs = [
        "The Agent class is the core component of the Strands framework.",
        "",
        "Constructor Parameters:",
        "  - model: The language model to use (BedrockModel, OpenAIModel, etc.)",
        "  - system_prompt: Instructions defining agent behavior",
        "  - tools: List of callable tools the agent can use",
        "  - memory: Memory system for context retention",
        "",
        "Methods:",
        "  - __call__(message: str) -> AgentResult: Send a message to the agent",
        "  - reset(): Clear conversation history",
        "  - add_tool(tool): Add a new tool to the agent",
        "",
        "Example:",
        "  from strands import Agent",
        "  agent = Agent(system_prompt='You are helpful')",
        "  response = agent('Hello!')"
    ]
    
    for line in agent_docs:
        if y_pos < 100:
            c.showPage()
            y_pos = height - 100
        c.drawString(100, y_pos, line)
        y_pos -= 20
    
    # Page 2 - Models
    c.showPage()
    c.setFont("Helvetica-Bold", 16)
    c.drawString(100, height - 100, "Model Classes")
    c.setFont("Helvetica", 11)
    y_pos = height - 130
    
    model_docs = [
        "BedrockModel:",
        "  - AWS Bedrock integration for Claude, Llama, and other models",
        "  - Requires AWS credentials and Bedrock access",
        "  - model_id: 'claude-3.7-sonnet', 'llama-3.1-8b', etc.",
        "",
        "OpenAIModel:",
        "  - OpenAI API integration for GPT models",
        "  - Requires OPENAI_API_KEY environment variable",
        "  - model_id: 'gpt-4', 'gpt-3.5-turbo', etc.",
        "",
        "OllamaModel:",
        "  - Local model execution with Ollama",
        "  - No API keys required, runs on your machine",
        "  - model_id: 'llama3.2', 'mistral', 'phi3', etc.",
        "",
        "All models support:",
        "  - Streaming responses",
        "  - Token counting",
        "  - Custom parameters"
    ]
    
    for line in model_docs:
        if y_pos < 100:
            c.showPage()
            y_pos = height - 100
        c.drawString(100, y_pos, line)
        y_pos -= 20
    
    # Page 3 - Tools
    c.showPage()
    c.setFont("Helvetica-Bold", 16)
    c.drawString(100, height - 100, "Creating Tools")
    c.setFont("Helvetica", 11)
    y_pos = height - 130
    
    tool_docs = [
        "Tools extend agent capabilities with custom functions.",
        "",
        "Creating a Tool:",
        "  from strands import tool",
        "  ",
        "  @tool",
        "  def calculate_sum(a: int, b: int) -> int:",
        '      """Add two numbers together."""',
        "      return a + b",
        "",
        "Tool Requirements:",
        "  - Must have type hints for all parameters",
        "  - Must have a docstring describing the function",
        "  - Should handle errors gracefully",
        "  - Return JSON-serializable data",
        "",
        "Advanced Tool Features:",
        "  - Async tool support",
        "  - Tool validation",
        "  - Custom error handling",
        "  - Tool composition"
    ]
    
    for line in tool_docs:
        if y_pos < 100:
            c.showPage()
            y_pos = height - 100
        c.drawString(100, y_pos, line)
        y_pos -= 20
    
    c.save()
    
    # Create DOCX - Agent Design Patterns
    if DOCX_AVAILABLE:
        doc = Document()
        
        # Title
        title = doc.add_heading('Strands Agent Design Patterns', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f'Version 0.1.x | Generated: {datetime.now().strftime("%Y-%m-%d")}')
        
        # Pattern 1: Research Agent
        doc.add_heading('Pattern 1: Research Agent', level=1)
        
        doc.add_paragraph(
            'The Research Agent pattern combines web search, document analysis, '
            'and fact-checking capabilities to provide comprehensive research assistance.'
        )
        
        doc.add_heading('Implementation', level=2)
        doc.add_paragraph('Key components:')
        doc.add_paragraph('• Web search tool for finding information', style='List Bullet')
        doc.add_paragraph('• Document parser for analyzing sources', style='List Bullet')
        doc.add_paragraph('• Fact-checking tool for verification', style='List Bullet')
        doc.add_paragraph('• Citation manager for references', style='List Bullet')
        
        # Code example
        doc.add_heading('Code Example', level=3)
        code_para = doc.add_paragraph()
        code_para.add_run('from strands import Agent, tool\n\n').font.name = 'Courier New'
        code_para.add_run('@tool\n').font.name = 'Courier New'
        code_para.add_run('def search_web(query: str) -> List[str]:\n').font.name = 'Courier New'
        code_para.add_run('    """Search the web for information."""\n').font.name = 'Courier New'
        code_para.add_run('    # Implementation here\n').font.name = 'Courier New'
        code_para.add_run('    pass\n\n').font.name = 'Courier New'
        code_para.add_run('research_agent = Agent(\n').font.name = 'Courier New'
        code_para.add_run('    system_prompt="You are a research assistant.",\n').font.name = 'Courier New'
        code_para.add_run('    tools=[search_web, parse_document, check_facts]\n').font.name = 'Courier New'
        code_para.add_run(')\n').font.name = 'Courier New'
        
        # Pattern 2: Customer Service Agent
        doc.add_page_break()
        doc.add_heading('Pattern 2: Customer Service Agent', level=1)
        
        doc.add_paragraph(
            'The Customer Service Agent pattern handles user inquiries, '
            'accesses knowledge bases, and escalates complex issues.'
        )
        
        doc.add_heading('Key Features', level=2)
        doc.add_paragraph('• FAQ search capability', style='List Bullet')
        doc.add_paragraph('• Order status lookup', style='List Bullet')
        doc.add_paragraph('• Ticket creation system', style='List Bullet')
        doc.add_paragraph('• Sentiment analysis', style='List Bullet')
        doc.add_paragraph('• Escalation workflow', style='List Bullet')
        
        doc.add_heading('Best Practices', level=2)
        doc.add_paragraph(
            '1. Always maintain a professional tone\n'
            '2. Search knowledge base before responding\n'
            '3. Collect necessary information before escalating\n'
            '4. Provide clear next steps to users\n'
            '5. Log all interactions for quality assurance'
        )
        
        # Pattern 3: Code Assistant
        doc.add_page_break()
        doc.add_heading('Pattern 3: Code Assistant', level=1)
        
        doc.add_paragraph(
            'The Code Assistant pattern helps developers write, debug, '
            'and optimize code across multiple programming languages.'
        )
        
        doc.add_heading('Core Components', level=2)
        doc.add_paragraph('• Code analysis tool', style='List Bullet')
        doc.add_paragraph('• Syntax checker', style='List Bullet')
        doc.add_paragraph('• Test generator', style='List Bullet')
        doc.add_paragraph('• Documentation writer', style='List Bullet')
        doc.add_paragraph('• Performance profiler', style='List Bullet')
        
        doc.add_heading('Implementation Strategy', level=2)
        doc.add_paragraph(
            'The Code Assistant should be able to:\n'
            '1. Understand code context and requirements\n'
            '2. Generate code following best practices\n'
            '3. Explain complex code segments\n'
            '4. Suggest optimizations and improvements\n'
            '5. Help with debugging and error resolution'
        )
        
        # Save the document
        docx_path = rag_docs_path / "agent_design_patterns.docx"
        doc.save(str(docx_path))
    
    # Create troubleshooting guide PDF
    trouble_pdf_path = rag_docs_path / "troubleshooting_guide.pdf"
    c2 = canvas.Canvas(str(trouble_pdf_path), pagesize=letter)
    
    # Title page
    c2.setFont("Helvetica-Bold", 20)
    c2.drawString(100, height - 100, "Strands Troubleshooting Guide")
    c2.setFont("Helvetica", 12)
    c2.drawString(100, height - 140, "Common Issues and Solutions")
    
    # Issue 1
    c2.setFont("Helvetica-Bold", 14)
    c2.drawString(100, height - 200, "Issue 1: Agent Not Responding")
    c2.setFont("Helvetica", 11)
    y_pos = height - 230
    
    solutions = [
        "Symptoms: Agent returns None or empty responses",
        "",
        "Solutions:",
        "1. Check model configuration:",
        "   - Verify API keys are set correctly",
        "   - Ensure model ID is valid",
        "   - Check network connectivity",
        "",
        "2. Validate input:",
        "   - Ensure message is not empty",
        "   - Check for special characters",
        "   - Verify encoding is UTF-8",
        "",
        "3. Review system prompt:",
        "   - Avoid contradictory instructions",
        "   - Keep prompts clear and concise",
        "   - Test with minimal prompt first"
    ]
    
    for line in solutions:
        if y_pos < 100:
            c2.showPage()
            y_pos = height - 100
        c2.drawString(100, y_pos, line)
        y_pos -= 20
    
    # Issue 2
    c2.showPage()
    c2.setFont("Helvetica-Bold", 14)
    c2.drawString(100, height - 100, "Issue 2: Tool Execution Failures")
    c2.setFont("Helvetica", 11)
    y_pos = height - 130
    
    tool_issues = [
        "Symptoms: Agent tries to use tools but they fail",
        "",
        "Common Causes:",
        "1. Missing type hints:",
        "   - All parameters must have type annotations",
        "   - Return type must be specified",
        "   - Use proper types (int, str, List, etc.)",
        "",
        "2. Invalid docstring:",
        "   - Tool must have a docstring",
        "   - Docstring should describe the function clearly",
        "   - Include parameter descriptions",
        "",
        "3. Error handling:",
        "   - Tools should handle exceptions",
        "   - Return meaningful error messages",
        "   - Don't let tools crash the agent",
        "",
        "Example of correct tool:",
        "@tool",
        "def calculate(x: int, y: int) -> int:",
        '    """Add two numbers together."""',
        "    return x + y"
    ]
    
    for line in tool_issues:
        if y_pos < 100:
            c2.showPage()
            y_pos = height - 100
        c2.drawString(100, y_pos, line)
        y_pos -= 20
    
    # Issue 3
    c2.showPage()
    c2.setFont("Helvetica-Bold", 14)
    c2.drawString(100, height - 100, "Issue 3: AWS Bedrock Access Denied")
    c2.setFont("Helvetica", 11)
    y_pos = height - 130
    
    bedrock_issues = [
        "Symptoms: 'Access Denied' or '403 Forbidden' errors",
        "",
        "Checklist:",
        "1. Model Access:",
        "   - Log into AWS Console",
        "   - Navigate to Bedrock > Model access",
        "   - Ensure Claude 3.7 Sonnet is enabled",
        "   - Wait for approval (can take minutes)",
        "",
        "2. IAM Permissions:",
        "   - User/role needs 'bedrock:InvokeModel'",
        "   - Check policy attachments",
        "   - Verify resource permissions",
        "",
        "3. Region Settings:",
        "   - Bedrock not available in all regions",
        "   - Check your AWS_DEFAULT_REGION",
        "   - Try us-west-2 or us-east-1",
        "",
        "4. Credentials:",
        "   - Run 'aws configure list' to verify",
        "   - Check environment variables",
        "   - Ensure credentials are not expired"
    ]
    
    for line in bedrock_issues:
        if y_pos < 100:
            c2.showPage()
            y_pos = height - 100
        c2.drawString(100, y_pos, line)
        y_pos -= 20
    
    # Performance optimization page
    c2.showPage()
    c2.setFont("Helvetica-Bold", 14)
    c2.drawString(100, height - 100, "Performance Optimization")
    c2.setFont("Helvetica", 11)
    y_pos = height - 130
    
    perf_tips = [
        "Tips for Faster Agents:",
        "",
        "1. Model Selection:",
        "   - Use smaller models for simple tasks",
        "   - Claude 3.7 Haiku for speed",
        "   - Local models for zero latency",
        "",
        "2. Streaming Responses:",
        "   - Enable streaming for long outputs",
        "   - Improves perceived performance",
        "   - Better user experience",
        "",
        "3. Caching:",
        "   - Cache common responses",
        "   - Store tool results when appropriate",
        "   - Use memory systems effectively",
        "",
        "4. Batch Processing:",
        "   - Group similar requests",
        "   - Process documents in batches",
        "   - Parallelize when possible"
    ]
    
    for line in perf_tips:
        if y_pos < 100:
            c2.showPage()
            y_pos = height - 100
        c2.drawString(100, y_pos, line)
        y_pos -= 20
    
    c2.save()
    
    # Create case study documents
    case_studies_path = rag_docs_path / "case_studies"
    case_studies_path.mkdir(exist_ok=True)
    
    # E-commerce chatbot case study
    ecommerce_content = """# E-commerce Chatbot Case Study

## Project Overview

**Client**: FashionFlow Online Retail  
**Challenge**: Handle 10,000+ daily customer inquiries with personalized responses  
**Solution**: Strands-powered AI chatbot with product search and order tracking

## Implementation Details

### Agent Architecture

The e-commerce chatbot uses a multi-tool approach:

1. **Product Search Tool**: Searches inventory database
2. **Order Tracking Tool**: Retrieves order status
3. **FAQ Tool**: Answers common questions
4. **Recommendation Tool**: Suggests products based on preferences

### System Prompt

```
You are FashionFlow's friendly shopping assistant. Help customers find products, 
track orders, and answer questions. Be enthusiastic about fashion and provide 
personalized recommendations. Always be helpful and professional.
```

### Key Features

- Natural language product search
- Size and fit recommendations
- Order status updates
- Return policy information
- Personalized styling suggestions

## Results

- **Response Time**: 90% of queries answered in <2 seconds
- **Customer Satisfaction**: 4.8/5 star rating
- **Cost Reduction**: 70% decrease in support tickets
- **Sales Impact**: 15% increase in conversion rate

## Lessons Learned

1. **Clear tool boundaries**: Each tool should have a specific purpose
2. **Fallback handling**: Always have a plan for edge cases
3. **Context retention**: Remember customer preferences within session
4. **Testing variety**: Test with diverse customer personas

## Code Snippet

```python
@tool
def search_products(query: str, category: str = None, max_price: float = None):
    \"\"\"Search for products in the inventory.\"\"\"
    # Connect to product database
    # Apply filters
    # Return top results
    pass

fashion_agent = Agent(
    model=bedrock_model,
    system_prompt=FASHION_PROMPT,
    tools=[search_products, track_order, get_faq, recommend_products]
)
```
"""
    
    with open(case_studies_path / "ecommerce_chatbot.md", "w") as f:
        f.write(ecommerce_content)
    
    # Create financial advisor PDF
    financial_pdf_path = case_studies_path / "financial_advisor_agent.pdf"
    c3 = canvas.Canvas(str(financial_pdf_path), pagesize=letter)
    
    c3.setFont("Helvetica-Bold", 18)
    c3.drawString(100, height - 100, "Financial Advisor Agent Case Study")
    
    c3.setFont("Helvetica", 11)
    y_pos = height - 140
    
    financial_content = [
        "Client: WealthWise Financial Services",
        "Challenge: Provide 24/7 financial guidance while ensuring compliance",
        "",
        "Solution Architecture:",
        "- Strands Agent with specialized financial tools",
        "- Integration with market data APIs",
        "- Compliance checking system",
        "- Risk assessment calculator",
        "",
        "Key Components:",
        "1. Market Analysis Tool: Real-time stock and fund data",
        "2. Portfolio Analyzer: Assess current holdings",
        "3. Risk Calculator: Determine risk tolerance",
        "4. Compliance Checker: Ensure advice meets regulations",
        "",
        "System Prompt:",
        "You are a licensed financial advisor. Provide educational",
        "information only. Always include disclaimers. Never guarantee",
        "returns. Encourage users to consult human advisors for",
        "personalized advice.",
        "",
        "Results:",
        "- 50,000+ users served monthly",
        "- 99.9% compliance rate",
        "- 4.7/5 user satisfaction",
        "- $2M cost savings annually"
    ]
    
    for line in financial_content:
        if y_pos < 100:
            c3.showPage()
            y_pos = height - 100
        c3.drawString(100, y_pos, line)
        y_pos -= 20
    
    c3.save()
    
    # Healthcare assistant case study
    healthcare_content = """# Healthcare Assistant Case Study

## Background

**Organization**: MediCare Plus Network  
**Objective**: Provide HIPAA-compliant health information and appointment scheduling  
**Constraints**: Strict privacy requirements, no diagnosis capability

## Solution Design

### Agent Configuration

```python
healthcare_agent = Agent(
    model=bedrock_model,
    system_prompt=\"\"\"You are a healthcare information assistant. 
    You can provide general health information and help with appointments.
    Never diagnose conditions or recommend treatments. Always advise 
    consulting healthcare professionals for medical concerns.\"\"\",
    tools=[search_symptoms, find_doctors, schedule_appointment, get_health_info]
)
```

### Privacy Measures

1. **No Personal Data Storage**: Agent doesn't retain PHI between sessions
2. **Encrypted Communications**: All data transmission encrypted
3. **Audit Logging**: All interactions logged for compliance
4. **Access Controls**: Role-based access to different features

## Implementation Highlights

### Symptom Checker Tool
- Provides general information only
- Always includes disclaimers
- Suggests when to seek immediate care

### Doctor Finder Tool
- Searches by specialty and location
- Shows availability and ratings
- Integrates with scheduling system

### Appointment Scheduler
- Real-time availability checking
- Automated reminders
- Rescheduling capabilities

## Outcomes

- **Usage**: 100,000+ monthly interactions
- **Appointment Booking**: 30% increase in online scheduling
- **Patient Satisfaction**: 4.9/5 rating
- **Compliance**: 100% HIPAA compliant
- **Cost Savings**: $500K annually in call center costs

## Key Learnings

1. **Clear Boundaries**: Agent must know what it cannot do
2. **Compliance First**: Every feature designed with HIPAA in mind
3. **User Education**: Clear about AI limitations
4. **Human Handoff**: Seamless escalation to human staff when needed
"""
    
    with open(case_studies_path / "healthcare_assistant.md", "w") as f:
        f.write(healthcare_content)
    
    # Return list of created files
    files_created = [
        "strands_api_reference.pdf",
        "agent_design_patterns.docx" if DOCX_AVAILABLE else None,
        "troubleshooting_guide.pdf",
        "case_studies/ecommerce_chatbot.md",
        "case_studies/financial_advisor_agent.pdf",
        "case_studies/healthcare_assistant.md"
    ]
    
    return [f for f in files_created if f is not None]

if __name__ == "__main__":
    # Create all documents
    print("Creating basic documents...")
    basic_files = create_basic_documents()
    print(f"Created {len(basic_files)} basic documents")
    
    print("\nCreating advanced documents...")
    advanced_files = create_advanced_documents()
    print(f"Created {len(advanced_files)} advanced documents")
    
    print("\nAll documents created successfully!")
